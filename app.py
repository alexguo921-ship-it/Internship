"""
AI内部竞聘评审系统 v3.0
腾讯AIHR面试作品 · 项目运营方视角
"""

import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
import io, json, re, uuid, os
from datetime import datetime, timedelta
from pathlib import Path
from anthropic import Anthropic
import data_manager as dm

try:
    import pdfplumber; HAS_PDF = True
except ImportError:
    HAS_PDF = False
try:
    from docx import Document as DocxDoc
    from docx.shared import Pt, RGBColor, Cm
    from docx.oxml.ns import qn as _qn
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
try:
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor, black
    HAS_REPORTLAB = True
except ImportError:
    HAS_REPORTLAB = False

# ═══════════════════════════════════════════════════════════════
# PAGE CONFIG
# ═══════════════════════════════════════════════════════════════
st.set_page_config(
    page_title="AI内部竞聘评审系统",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.markdown("""
<style>
/* ── Base ─────────────────────────────────────────────────── */
html, body, [class*="css"] {
    font-family: -apple-system, BlinkMacSystemFont, "PingFang SC",
                 "Microsoft YaHei", "Segoe UI", sans-serif;
}
.stApp { background-color: #f4f6fb; }

/* ── Typography ───────────────────────────────────────────── */
p, li, .stMarkdown p, .stMarkdown li {
    font-size: 14px !important;
    line-height: 1.9 !important;
    color: #262626;
}
h1 { font-size: 22px !important; font-weight: 700 !important;
     letter-spacing: -0.3px; color: #141414 !important; }
h2 { font-size: 18px !important; font-weight: 600 !important; color: #1f1f1f !important; }
h3 { font-size: 15px !important; font-weight: 600 !important; color: #333 !important; }
h4 { font-size: 13.5px !important; font-weight: 600 !important; }

/* ── Metric cards ─────────────────────────────────────────── */
div[data-testid="metric-container"] {
    background: white;
    border: 1px solid #e6eaf0;
    border-radius: 12px;
    padding: 18px 22px;
    box-shadow: 0 1px 6px rgba(0,0,0,0.06);
    transition: box-shadow .2s;
}
div[data-testid="metric-container"]:hover {
    box-shadow: 0 3px 12px rgba(0,0,0,0.10);
}
[data-testid="stMetricValue"] { font-size: 28px !important; font-weight: 700 !important; }
[data-testid="stMetricLabel"] { font-size: 13px !important; color: #595959 !important; }

/* ── Cards ────────────────────────────────────────────────── */
.card {
    background: white; border: 1px solid #e6eaf0;
    border-radius: 12px; padding: 22px 26px;
    box-shadow: 0 1px 6px rgba(0,0,0,0.06); margin-bottom: 14px;
}
.card-blue  { border-top: 3px solid #1890ff; }
.card-green { border-top: 3px solid #52c41a; }
.card-gold  { border-top: 3px solid #faad14; }

/* ── Report box ───────────────────────────────────────────── */
.report-box {
    background: white; border: 1px solid #e6eaf0;
    border-radius: 12px; padding: 32px 36px;
    line-height: 2.0; font-size: 14px;
    box-shadow: 0 1px 8px rgba(0,0,0,0.06);
}
.report-box h3 { margin-top: 20px; margin-bottom: 8px; }
.report-box li { margin-bottom: 6px; }

/* ── Badges ───────────────────────────────────────────────── */
.badge-green  { background:#f6ffed; color:#389e0d; border:1px solid #b7eb8f;
                padding:3px 12px; border-radius:20px; font-size:12px; font-weight:600; }
.badge-yellow { background:#fffbe6; color:#d48806; border:1px solid #ffe58f;
                padding:3px 12px; border-radius:20px; font-size:12px; font-weight:600; }
.badge-red    { background:#fff2f0; color:#cf1322; border:1px solid #ffa39e;
                padding:3px 12px; border-radius:20px; font-size:12px; font-weight:600; }

/* ── Pulse dot ────────────────────────────────────────────── */
.pulse {
    display:inline-block; width:10px; height:10px;
    background:#52c41a; border-radius:50%;
    animation: pulse 1.8s ease-in-out infinite;
    margin-right: 6px; vertical-align: middle;
}
@keyframes pulse {
    0%,100% { box-shadow: 0 0 0 0 rgba(82,196,26,.5); }
    50% { box-shadow: 0 0 0 7px rgba(82,196,26,0); }
}

/* ── Tabs ─────────────────────────────────────────────────── */
.stTabs [data-baseweb="tab-list"] { gap: 4px; }
.stTabs [data-baseweb="tab"] {
    font-size: 13px !important; font-weight: 500;
    padding: 10px 18px !important; border-radius: 8px 8px 0 0;
}
button[data-testid="stBaseButton-secondary"] { font-size: 13px !important; }
button[data-testid="stBaseButton-primary"]   { font-size: 13px !important; }

/* ── Portal banner ────────────────────────────────────────── */
.portal-banner {
    background: linear-gradient(135deg,#0050b3,#1890ff);
    color: white; padding: 28px 36px; border-radius: 14px;
    margin-bottom: 22px; line-height: 1.7;
}
.portal-banner h2 { color: white !important; font-size: 20px !important; margin: 0 0 6px; }
.portal-banner p  { margin: 0; opacity: .88; font-size: 13.5px; }

/* ── Monitor live chip ────────────────────────────────────── */
.live-chip {
    display:inline-flex; align-items:center;
    background:#f6ffed; border:1px solid #b7eb8f;
    color:#389e0d; border-radius:20px;
    padding:4px 14px; font-size:13px; font-weight:600; gap:6px;
}
/* ── Info box ─────────────────────────────────────────────── */
.info-box {
    background: #e6f4ff; border: 1px solid #91caff;
    border-radius: 10px; padding: 14px 18px;
    font-size: 13.5px; line-height: 1.8; color: #003eb3;
}
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
# CONSTANTS
# ═══════════════════════════════════════════════════════════════
SURNAMES = list("张李王刘陈杨黄赵周吴徐孙马朱胡林郭何高罗")
GIVEN    = ["伟","芳","娜","敏","静","丽","强","磊","洋","艳","勇","军","杰","涛",
            "明","超","霞","平","刚","玲","华","辉","飞","云","红","建","宇","鑫","斌","梅"]
DEPTS    = ["产品部","技术部","运营部","市场部","数据部","设计部","商业化部","用研部"]
DEPT_P   = [0.10,0.22,0.18,0.14,0.13,0.08,0.08,0.07]
PAGES    = ["总览仪表板","候选人门户","能力发展报告","申诉通道","员工信息","人工复核"]
PAGE_W   = [0.25, 0.30, 0.20, 0.12, 0.08, 0.05]

RESOURCES = {
    "岗位匹配度": [
        ("腾讯内部·PM成长路径手册（必读）",        "https://iwiki.woa.com/p/pm-growth-path"),
        ("腾讯课堂·产品经理核心技能训练营（6h）",    "https://ke.qq.com/course/5982909"),
        ("极客时间·产品思维与产品管理实战专栏",      "https://time.geekbang.org/column/intro/100003101"),
        ("人人都是产品经理·需求分析与功能设计指南",   "https://www.woshipm.com/pmd/5847200.html"),
    ],
    "历史绩效": [
        ("腾讯内部·OKR设定与绩效复盘方法论",        "https://km.woa.com/articles/show/okr-guide"),
        ("腾讯内部·绩效提升工作坊（每季度开放报名）", "https://talent.tencent.com/training/performance"),
        ("极客时间·高效工作法与个人效能提升专栏",    "https://time.geekbang.org/column/intro/100018401"),
    ],
    "笔试成绩": [
        ("腾讯内部·PM笔试历年题库与详解",           "https://km.woa.com/group/product/exam-prep"),
        ("极客时间·数据分析实战 45 讲",             "https://time.geekbang.org/column/intro/100093501"),
        ("产品经理之家·案例题在线练习库",            "https://www.pmcaff.com/tools/exam"),
        ("腾讯数据平台·数据思维与分析入门课",        "https://data.qq.com/article/data-thinking"),
    ],
}

APPEAL_CATS = ["AI评分偏差","绩效数据滞后","笔试题目歧义","材料未被识别","其他"]

# ═══════════════════════════════════════════════════════════════
# ANALYTICS SEED DATA
# ═══════════════════════════════════════════════════════════════
@st.cache_data
def generate_analytics_seed():
    rng = np.random.default_rng(88)
    hour_w = np.array([0.3,0.2,0.1,0.1,0.1,0.2,0.5,1.5,3.2,4.1,
                       3.6,2.6,2.1,1.6,3.6,4.2,3.1,2.6,2.1,1.6,
                       1.1,0.9,0.7,0.5])
    hour_w = hour_w / hour_w.sum()
    base = datetime.now() - timedelta(days=30)
    visits = []
    for day in range(30):
        d = base + timedelta(days=day)
        mult = 1.0 if d.weekday() < 5 else 0.45
        n = int(rng.integers(35, 95) * mult)
        for i in range(n):
            h = int(rng.choice(24, p=hour_w))
            m = int(rng.integers(0, 60))
            ts = d.replace(hour=h, minute=m, second=0, microsecond=0)
            visits.append({
                "ts": ts,
                "page": rng.choice(PAGES, p=PAGE_W),
                "sid": f"seed_{day*200+i:05d}",
            })
    return visits

@st.cache_data
def generate_kpi_seed():
    rng = np.random.default_rng(77)
    dates = pd.date_range(end=datetime.now(), periods=30, freq='D')
    before_days  = rng.integers(25, 36, 30).astype(float)
    after_days   = np.clip(rng.normal(12.2, 1.2, 30), 9, 16)
    before_comp  = rng.uniform(18, 25, 30)
    after_comp   = np.clip(rng.normal(3.8, 0.7, 30), 1.5, 6.5)
    before_rate  = rng.uniform(25, 32, 30)
    after_rate   = np.clip(rng.normal(13.2, 1.1, 30), 10, 16)
    return pd.DataFrame({
        "日期":          [d.strftime("%Y-%m-%d") for d in dates],
        "评审时长_AI前(天)":  np.round(before_days, 0).astype(int),
        "评审时长_AI后(天)":  np.round(after_days, 1),
        "投诉率_AI前(%)":    np.round(before_comp, 1),
        "投诉率_AI后(%)":    np.round(after_comp, 1),
        "人工复核率_AI后(%)": np.round(after_rate, 1),
        "1年离职率_AI后(%)": np.round(np.clip(rng.normal(7.5, 0.8, 30), 5, 10), 1),
    })

# ═══════════════════════════════════════════════════════════════
# DATA GENERATION
# ═══════════════════════════════════════════════════════════════
@st.cache_data
def _gen_base():
    rng = np.random.default_rng(42)
    n   = 360
    lat = rng.normal(0, 1, n)
    pm  = np.clip(60 + 10*lat + 10*rng.normal(0,1,n), 0, 100)
    pe  = np.clip(65 +  8*lat +  7*rng.normal(0,1,n), 0, 100)
    wt  = np.clip(58 +  6*lat + 14*rng.normal(0,1,n), 0, 100)
    r2  = np.random.default_rng(123)
    names = [SURNAMES[r2.integers(0,20)] + GIVEN[r2.integers(0,30)] for _ in range(n)]
    depts = [DEPTS[i] for i in r2.choice(len(DEPTS), n, p=DEPT_P)]
    years = r2.uniform(1.0, 12.0, n)
    return pd.DataFrame({
        "候选人":     names,
        "员工编号":   [f"EMP{10001+i}" for i in range(n)],
        "当前部门":   depts,
        "工龄(年)":   np.round(years, 1),
        "岗位匹配度":  np.round(pm, 1),
        "历史绩效":    np.round(pe, 1),
        "笔试成绩":    np.round(wt, 1),
        "来源":        ["系统导入"] * n,
    })

# ═══════════════════════════════════════════════════════════════
# CORE HELPERS
# ═══════════════════════════════════════════════════════════════
def apply_weights(base, wm, wp, wt_, g, r):
    df = base.copy()
    df["AI综合评分"] = (df["岗位匹配度"]*wm + df["历史绩效"]*wp + df["笔试成绩"]*wt_).round(1)
    Z = {"绿区":"✅ AI通过","黄区":"⏳ 待人工复核","红区":"❌ AI淘汰"}
    df["评审区间"] = df["AI综合评分"].apply(
        lambda s: "绿区" if s>=g else ("黄区" if s>=r else "红区"))
    df["状态"] = df["评审区间"].map(Z)
    df = df.sort_values("AI综合评分", ascending=False).reset_index(drop=True)
    df["排名"] = df.index + 1
    return df

def extract_text(f):
    raw = f.getvalue(); name = f.name.lower()
    if name.endswith(".pdf") and HAS_PDF:
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            return "\n".join(p.extract_text() or "" for p in pdf.pages)
    if name.endswith(".docx") and HAS_DOCX:
        return "\n".join(p.text for p in DocxDoc(io.BytesIO(raw)).paragraphs)
    if name.endswith(".txt"):
        return raw.decode("utf-8", errors="ignore")
    return ""

def heuristic_parse(text, fname):
    nm = re.search(r'姓名[：:]\s*(\S+)', text)
    yr = re.search(r'(\d+)\s*年.*?工作', text)
    return {"name": nm.group(1) if nm else fname.rsplit(".",1)[0],
            "department":"待补充","years":float(yr.group(1)) if yr else 3.0,
            "education":"待补充","skills":[],
            "position_match_estimate":60.0,"performance_estimate":65.0,
            "written_test_estimate":58.0,"highlights":"请手动补充亮点"}

def ai_parse(client, text, fname):
    try:
        r = client.messages.create(model="claude-sonnet-4-6", max_tokens=500,
            messages=[{"role":"user","content":
                f'从以下简历提取信息，返回纯JSON：\n{text[:2500]}\n格式：'
                '{"name":"","department":"","years":0,"education":"","skills":[],'
                '"position_match_estimate":0,"performance_estimate":0,'
                '"written_test_estimate":0,"highlights":""}'}])
        m = re.search(r'\{.*\}', r.content[0].text, re.DOTALL)
        return json.loads(m.group()) if m else {}
    except:
        return {}

# ═══════════════════════════════════════════════════════════════
# APPEAL CLASSIFIER
# ═══════════════════════════════════════════════════════════════
def classify_appeal_rules(reason: str) -> dict:
    t = reason.lower()
    if any(w in t for w in ["项目","经历","未被识别","评分偏","不准确","ai识别","识别错误"]):
        cat, urg = "AI评分偏差", "高"
    elif any(w in t for w in ["绩效","数据","更新","滞后","系统未","kpi"]):
        cat, urg = "绩效数据滞后", "中"
    elif any(w in t for w in ["题目","笔试","歧义","理解","题目表述"]):
        cat, urg = "笔试题目歧义", "中"
    elif any(w in t for w in ["材料","附件","证明","证书","未上传"]):
        cat, urg = "材料未被识别", "高"
    else:
        cat, urg = "其他", "低"
    claims = [s.strip() for s in re.split(r'[，。！？；]', reason) if len(s.strip()) > 8][:3]
    resp_map = {
        "AI评分偏差":   "您好，已收到您关于AI评分偏差的申诉，HR将重点复核您的项目经历描述与岗位能力映射，请保持联系方式畅通。",
        "绩效数据滞后": "您好，已收到您关于绩效数据的申诉，HR将联系绩效管理团队核实最新数据，预计1个工作日内完成数据校验。",
        "笔试题目歧义": "您好，已收到您关于笔试题目的申诉，题目审核小组将在1个工作日内进行专项复核并给出书面回复。",
        "材料未被识别": "您好，已收到您关于材料识别的申诉，请通过邮件重新提交相关材料至 hr-ai@tencent.com，将加急处理。",
        "其他":         "您好，已收到您的申诉，HR将在3个工作日内完成人工复核并给出书面回复。",
    }
    return {"category": cat, "urgency": urg,
            "key_claims": claims, "auto_response": resp_map[cat]}

def ai_classify_appeal(client, reason: str) -> dict:
    try:
        r = client.messages.create(model="claude-sonnet-4-6", max_tokens=400,
            messages=[{"role":"user","content":
                f'对以下竞聘申诉分类，返回纯JSON：\n{reason}\n格式：'
                '{"category":"AI评分偏差|绩效数据滞后|笔试题目歧义|材料未被识别|其他",'
                '"urgency":"高|中|低","key_claims":["诉求1"],'
                '"auto_response":"给申诉人的温和回复50字以内","hr_focus":"HR核查重点"}'}])
        m = re.search(r'\{.*\}', r.content[0].text, re.DOTALL)
        return json.loads(m.group()) if m else classify_appeal_rules(reason)
    except:
        return classify_appeal_rules(reason)

# ═══════════════════════════════════════════════════════════════
# REPORT BUILDER
# ═══════════════════════════════════════════════════════════════
def build_report(row, df_all, g_thr, wm, wp, wt_):
    name  = row["候选人"]; dept = row["当前部门"]
    yrs   = row["工龄(年)"]; score = row["AI综合评分"]
    pm    = row["岗位匹配度"]; pe = row["历史绩效"]; wt = row["笔试成绩"]
    gap   = g_thr - score; now = datetime.now().strftime("%Y-%m-%d %H:%M")

    def pct(col, val): return (df_all[col] <= val).mean() * 100

    dims = [("岗位匹配度",pm,pct("岗位匹配度",pm),wm),
            ("历史绩效",pe,pct("历史绩效",pe),wp),
            ("笔试成绩",wt,pct("笔试成绩",wt),wt_)]
    lo = min(dims, key=lambda x: x[1])
    hi = max(dims, key=lambda x: x[1])

    gap_desc = {
        "岗位匹配度": "产品视角与PM岗位契合度有待提升，需加强需求分析框架、用户研究方法和产品方法论的系统掌握",
        "历史绩效":   "历史项目可量化成果的呈现与OKR完成度记录需进一步完善，建议与直属上级重新梳理贡献数据",
        "笔试成绩":   "PM专业知识的结构化表达与案例推导能力尚需强化，可通过系统刷题和案例练习在短期内快速提升",
    }
    res_md = "\n".join(f"  - [{t}]({u})" for t, u in RESOURCES.get(lo[0],[]))

    return f"""## {name} · 能力发展报告

**评审日期**：{now}　　**竞聘岗位**：产品经理　　**当前部门**：{dept}　　**工龄**：{yrs} 年

---

### 一、综合评价

感谢您积极参与本次产品经理内部竞聘。您的 AI 综合评分为 **{score:.1f} 分**，超越了全体 {len(df_all)} 名报名员工中 **{pct("AI综合评分",score):.0f}%** 的候选人，与通过线（{g_thr} 分）仅差 **{gap:.1f} 分**——差距并不遥远。本报告将为您精准定位短板，并提供一条可立即执行的提升路径。

---

### 二、能力亮点

**▶ {hi[0]}（{hi[1]:.1f} 分 · 权重 {hi[3]:.0%} · 超越全体 {hi[2]:.0f}%）**
得分高于全体均值 {df_all[hi[0]].mean():.1f} 分，处于报名者前 {100-hi[2]:.0f}% 梯队，是您当前最显著的竞争优势。建议在下次竞聘材料中将相关项目成果量化呈现（如：主导 XX 功能上线，DAU 提升 X%）。

**▶ 内部跨团队协作经验（工龄 {yrs} 年）**
{yrs} 年的内部工作经验意味着您已深度熟悉公司文化、协作机制与业务背景，转岗后适应成本极低——这是外部候选人无法复制的结构性优势，在人工复核环节可作为重要加分项重点陈述。

**▶ 主动进取意识**
在 {len(df_all)} 名报名者中，能够主动迈出参与竞聘这一步，证明了您突破现有边界的职业成长驱动力——这是 PM 团队最看重的软性特质之一。

---

### 三、成长重点

**▶ {lo[0]}（{lo[1]:.1f} 分 · 权重 {lo[3]:.0%} · 超越全体 {lo[2]:.0f}%）——⚡ 优先突破项**
与全体均值（{df_all[lo[0]].mean():.1f} 分）差距 **{df_all[lo[0]].mean()-lo[1]:.1f} 分**，是本次评分未能越过通过线的核心原因。{gap_desc[lo[0]]}。

**量化目标**：将该维度提升 **{gap*1.3:.0f}–{gap*1.6:.0f} 分**，即可在下次竞聘中显著提高通过概率。

**▶ 笔试专项（{wt:.1f} 分）**
笔试考察产品案例推理、数据分析与结构化表达能力。经验表明，通过专项练习，**3–4 周内**可看到 5–10 分的明显进步。

---

### 四、🎯 行动建议与学习资源

#### 第一步：{lo[0]} 专项突破（本月，预计提升 {gap*1.3:.0f}+ 分）

{res_md}

> 💡 **执行建议**：选取上方 1–2 门资源系统学习（6–12 小时），同时每周输出 1 篇 500 字产品分析，既锻炼能力，也为下次竞聘积累补充材料。

#### 第二步：实践积累（第 2–4 周）

- 主动申请参与产品部门虚拟项目组：[腾讯内部·开放项目列表](https://km.woa.com/group/product-open-projects)
- 在现部门承接需求对接角色，用 STAR 格式记录可量化案例（「主导 XX 需求落地，指标提升 X%」）
- 结构化复盘每次跨部门协作经历：[协作复盘记录模板](https://iwiki.woa.com/p/collab-review-template)

#### 第三步：系统认证（第 5–12 周）

- 完成腾讯 PM 成长认证：[腾讯产品经理认证体系](https://talent.tencent.com/certification/pm)（结业证书可附于下次竞聘材料）
- 建立个人 PM 知识框架：[知识地图参考模板](https://iwiki.woa.com/p/pm-knowledge-map)
- 模拟笔试演练（建议竞聘前 4 周完成）：[PM 能力在线模拟考](https://talent.tencent.com/mock-test/pm)

#### 下次竞聘 Checklist

- [ ] **{lo[0]}** 提升至 {lo[1]+gap*1.5:.0f} 分以上
- [ ] 积累至少 **2 个**可量化的产品实践案例（STAR 格式）
- [ ] 完成腾讯 PM 认证课程并取得结业证明
- [ ] 提前 2 周确认绩效数据已更新至 HR 系统
- [ ] 准备 3 分钟结构化自我介绍（聚焦与 PM 岗的能力匹配）

---

### 五、发展寄语

{gap:.0f} 分的差距，是方向，不是终点。您今天种下的每一份努力，都将在下次竞聘中成为真实可见的优势。组织珍视的不只是最终通过的 50 人，更珍视每一位主动求变、勇于突破的成长者——期待在下次竞聘中看到更好的您。

---

> 📣 **对评分有异议？** 结果公示后 **48 小时内**可通过「候选人门户」或「申诉通道」提交补充材料，HR 将在 **3 个工作日**内书面回复。
>
> *由腾讯 AI 竞聘评审系统生成 · {now} · 咨询请联系 hr-ai@tencent.com*"""

# ═══════════════════════════════════════════════════════════════
# PDF & WORD EXPORT
# ═══════════════════════════════════════════════════════════════
def to_pdf_bytes(md):
    if not HAS_REPORTLAB: return None
    try: pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
    except: pass
    BLUE=HexColor('#1890ff'); GRAY=HexColor('#595959'); LGRAY=HexColor('#8c8c8c')
    sN=ParagraphStyle('n',fontName='STSong-Light',fontSize=10.5,leading=20)
    sH1=ParagraphStyle('h1',fontName='STSong-Light',fontSize=15,leading=26,
                        textColor=BLUE,spaceBefore=4,spaceAfter=4)
    sH2=ParagraphStyle('h2',fontName='STSong-Light',fontSize=12.5,leading=22,
                        textColor=HexColor('#222'),spaceBefore=10,spaceAfter=4)
    sH3=ParagraphStyle('h3',fontName='STSong-Light',fontSize=11,leading=20,
                        textColor=HexColor('#13c2c2'),spaceBefore=6)
    sB=ParagraphStyle('b',fontName='STSong-Light',fontSize=10,leading=18,leftIndent=14)
    sQ=ParagraphStyle('q',fontName='STSong-Light',fontSize=9.5,leading=17,
                       leftIndent=20,textColor=GRAY)
    sC=ParagraphStyle('c',fontName='STSong-Light',fontSize=9,leading=15,textColor=LGRAY)
    def cl(s):
        s=re.sub(r'\*\*(.+?)\*\*',r'\1',s); s=re.sub(r'\[(.+?)\]\(.+?\)',r'\1',s)
        return s.strip()
    story=[]
    for line in md.split('\n'):
        s=line.strip()
        if not s: story.append(Spacer(1,.2*cm)); continue
        if s=='---': story.append(HRFlowable(width="100%",thickness=.5,
                                              color=HexColor('#d9d9d9'),spaceAfter=4)); continue
        c=cl(s)
        if s.startswith('## '):   story+=[Paragraph(c,sH1),HRFlowable(width="100%",thickness=1.5,color=BLUE,spaceAfter=6)]
        elif s.startswith('### '): story.append(Paragraph(c,sH2))
        elif s.startswith('#### '): story.append(Paragraph(c,sH3))
        elif s.startswith('  - '): story.append(Paragraph('  · '+cl(s[4:]),sB))
        elif s.startswith('- [ ]'): story.append(Paragraph('☐ '+cl(s[5:]),sB))
        elif s.startswith('- ') or s.startswith('▶'): story.append(Paragraph('• '+cl(s.lstrip('-▶ ')),sB))
        elif s.startswith('> '): story.append(Paragraph(cl(s[2:]),sQ))
        elif s.startswith('*') and s.endswith('*'): story.append(Paragraph(cl(s),sC))
        else: story.append(Paragraph(c,sN))
    buf=io.BytesIO()
    SimpleDocTemplate(buf,pagesize=A4,rightMargin=2.5*cm,leftMargin=2.5*cm,
                      topMargin=2.5*cm,bottomMargin=2*cm).build(story)
    buf.seek(0); return buf.getvalue()

def to_word_bytes(md):
    if not HAS_DOCX: return None
    doc=DocxDoc()
    sty=doc.styles['Normal']; sty.font.name='微软雅黑'
    sty.element.rPr.rFonts.set(_qn('w:eastAsia'),'微软雅黑')
    def ar(para,text,bold=False,color=None,size=10.5):
        r=para.add_run(text); r.font.name='微软雅黑'; r.font.size=Pt(size)
        r.element.rPr.rFonts.set(_qn('w:eastAsia'),'微软雅黑')
        if bold: r.font.bold=True
        if color: r.font.color.rgb=RGBColor(*color)
        return r
    def cl(s):
        s=re.sub(r'\*\*(.+?)\*\*',r'\1',s)
        s=re.sub(r'\[(.+?)\]\((.+?)\)',r'\1 → \2',s); return s.strip()
    for line in md.split('\n'):
        s=line.strip()
        if not s: doc.add_paragraph(); continue
        if s=='---': doc.add_paragraph('─'*44); continue
        c=cl(s)
        if s.startswith('## '):    p=doc.add_heading(level=1); ar(p,s[3:],bold=True,color=(0x18,0x90,0xFF),size=15)
        elif s.startswith('### '): p=doc.add_heading(level=2); ar(p,s[4:],bold=True,size=12)
        elif s.startswith('#### '): p=doc.add_heading(level=3); ar(p,s[5:],bold=True,color=(0x13,0xC2,0xC2),size=11)
        elif s.startswith('  - '): p=doc.add_paragraph(style='List Bullet 2'); ar(p,cl(s[4:]))
        elif s.startswith('- [ ]'): p=doc.add_paragraph(style='List Bullet'); ar(p,'☐ '+cl(s[5:]))
        elif s.startswith('- ') or s.startswith('▶'): p=doc.add_paragraph(style='List Bullet'); ar(p,cl(s.lstrip('-▶ ')))
        elif s.startswith('> '): p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(1); ar(p,cl(s[2:]),color=(0x59,0x59,0x59))
        else: p=doc.add_paragraph(); ar(p,c)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.getvalue()

# ═══════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════
DEFAULTS = {
    "base_df": None, "human_reviews": {}, "generated_reports": {},
    "appeals": {}, "client": None, "upload_parsed": [],
    "kpi_df": None, "live_visits": [],
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

if st.session_state.base_df is None:
    st.session_state.base_df = _gen_base()
    dm.save_candidates(st.session_state.base_df)   # 同步到外部门户共享文件
if st.session_state.kpi_df is None:
    st.session_state.kpi_df = generate_kpi_seed()

# Track this session visit
if "sid" not in st.session_state:
    st.session_state.sid = str(uuid.uuid4())[:8]
    st.session_state.live_visits.append({
        "ts": datetime.now(), "page": "总览仪表板", "sid": st.session_state.sid
    })

seed_visits = generate_analytics_seed()

def get_analytics():
    all_v = seed_visits + st.session_state.live_visits
    return all_v

def get_active_users(minutes=15):
    """只统计真实 session，不伪造数字。"""
    cutoff = datetime.now() - timedelta(minutes=minutes)
    recent = [v for v in st.session_state.live_visits if v["ts"] > cutoff]
    return len(set(v["sid"] for v in recent))

# ═══════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════
with st.sidebar:
    st.markdown("## ⚙️ 系统配置")
    st.divider()
    st.markdown("**评分权重**")
    wm  = st.slider("岗位匹配度", 0.20, 0.60, 0.40, 0.05)
    wp  = st.slider("历史绩效",   0.10, 0.50, 0.30, 0.05)
    wt_ = st.slider("笔试成绩",   0.10, 0.50, 0.30, 0.05)
    tot = wm+wp+wt_
    if abs(tot-1.0) > 0.01:
        st.warning(f"权重合计 {tot:.2f}，已归一化")
        wm/=tot; wp/=tot; wt_/=tot
    fig_w=go.Figure(go.Bar(x=[wm,wp,wt_],y=["岗位匹配","历史绩效","笔试成绩"],
                            orientation="h",marker_color=["#1890ff","#52c41a","#722ed1"],
                            text=[f"{v:.0%}" for v in [wm,wp,wt_]],textposition="inside"))
    fig_w.update_layout(height=115,margin=dict(l=0,r=0,t=4,b=4),showlegend=False,
                         plot_bgcolor="rgba(0,0,0,0)",
                         xaxis=dict(showticklabels=False,showgrid=False),
                         yaxis=dict(showgrid=False))
    st.plotly_chart(fig_w,use_container_width=True)
    st.divider()
    st.markdown("**复核阈值**")
    g_thr=int(st.number_input("绿区下限",65,85,73))
    r_thr=int(st.number_input("红区上限",50,70,67))
    st.divider()
    st.markdown("**Claude API**")
    api_key=st.text_input("API Key",type="password",placeholder="sk-ant-...")
    if api_key:
        try: st.session_state.client=Anthropic(api_key=api_key); st.success("✅ API 已连接")
        except Exception as e: st.error(str(e))
    else: st.caption("未配置 → 使用智能模板")
    st.divider()
    n_total=len(st.session_state.base_df)
    n_upload=(st.session_state.base_df["来源"]=="简历上传").sum()
    st.caption(f"AI竞聘评审系统 v3.0")
    st.caption(f"候选人：{n_total} 人（含上传 {n_upload} 人）")

# ═══════════════════════════════════════════════════════════════
# LIVE DF
# ═══════════════════════════════════════════════════════════════
df       = apply_weights(st.session_state.base_df, wm, wp, wt_, g_thr, r_thr)
green_n  = (df["评审区间"]=="绿区").sum()
yellow_n = (df["评审区间"]=="黄区").sum()
red_n    = (df["评审区间"]=="红区").sum()
ALL_NAMES= df["候选人"].tolist()

# ═══════════════════════════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════════════════════════
st.markdown(f"""
<div style="background:linear-gradient(135deg,#1890ff,#722ed1);
  color:white;padding:22px 32px;border-radius:14px;margin-bottom:22px;
  box-shadow:0 4px 16px rgba(24,144,255,.25);">
  <h1 style="margin:0;font-size:24px;font-weight:700;color:white!important;">
    🎯 AI内部竞聘评审系统</h1>
  <p style="margin:7px 0 0;opacity:.88;font-size:13.5px;line-height:1.6;">
    产品经理岗位 · 50个名额 · {len(df)} 名候选人 · 两周双轨制高效评审
  </p>
</div>""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════
# TABS
# ═══════════════════════════════════════════════════════════════
tab1,tab2,tab3,tab4,tab5,tab6,tab7,tab8 = st.tabs([
    "📊 总览 & 监控",
    "🌐 候选人门户",
    "📤 简历导入",
    "👤 员工信息与对比",
    "🔍 人工复核",
    "📋 能力发展报告",
    "📣 申诉管理",
    "📈 效果评估",
])

# ════════════════════════════════════════════════════════════════
# TAB 1 · 总览 & 监控
# ════════════════════════════════════════════════════════════════
with tab1:
    # ── External portal stats (real data from analytics.json) ─
    portal_stats = dm.get_portal_stats(active_minutes=15)
    ext_active   = portal_stats["active"]    # 外部门户：15分钟内活跃用户
    ext_total    = portal_stats["total"]     # 外部门户：累计访问
    ext_today    = portal_stats["today"]     # 外部门户：今日访问
    ext_visits   = portal_stats["visits"]   # 外部门户：全量访问记录（供图表用）

    # 历史图表也混入种子数据（仅演示用）
    all_v_chart  = generate_analytics_seed() + [
        {**v, "ts": datetime.fromisoformat(v["ts"])} for v in ext_visits
    ]

    # 申诉总数：合并内部 + 外部门户申诉
    ext_appeals   = dm.load_appeals()
    merged_appeals = {**ext_appeals, **st.session_state.appeals}

    # ── External portal link banner ───────────────────────────
    st.markdown("""
    <div style="background:linear-gradient(135deg,#0050b3 0%,#1890ff 60%,#36cfc9 100%);
         color:white;padding:18px 26px;border-radius:12px;margin-bottom:16px;
         display:flex;align-items:center;gap:20px;">
      <div style="font-size:32px;line-height:1">🌐</div>
      <div style="flex:1">
        <div style="font-size:15px;font-weight:700;margin-bottom:4px;">
          候选人外部门户（对外开放）
        </div>
        <div style="font-size:13px;opacity:.9;line-height:1.6;">
          候选人通过以下地址自助查询报告、提交申诉 — 下方监控数据仅统计此外部渠道流量
        </div>
      </div>
      <div style="background:rgba(255,255,255,.18);border-radius:10px;
           padding:10px 18px;font-size:14px;font-weight:700;
           border:1px solid rgba(255,255,255,.35);white-space:nowrap;">
        🔗 http://localhost:8504
      </div>
    </div>
    """, unsafe_allow_html=True)

    # ── KPI row ──────────────────────────────────────────────
    k1,k2,k3,k4,k5,k6 = st.columns(6)
    k1.metric("报名总人数", str(len(df)))
    k2.metric("开放名额", "50")
    k3.metric("✅ AI通过",  str(green_n),  f"{green_n/len(df):.1%}")
    k4.metric("⏳ 待复核",  str(yellow_n), f"{yellow_n/len(df):.1%}")
    rev_rate=yellow_n/len(df)
    k5.metric("人工复核率", f"{rev_rate:.1%}",
              delta="达标 ✓" if rev_rate<=.15 else "超标 ✗",
              delta_color="normal" if rev_rate<=.15 else "inverse")
    k6.metric("📣 申诉总数", str(len(merged_appeals)), help="含外部门户+内部录入")

    st.divider()

    # ── Live monitoring bar（外部门户实时数据）────────────────
    st.markdown(
        '<div style="font-size:12px;color:#8c8c8c;margin-bottom:8px;">'
        '📡 以下实时指标仅统计 <b>外部候选人门户</b>（port 8504）的真实访问，不含内部系统流量</div>',
        unsafe_allow_html=True)
    col_live = st.columns([1,1,1,1])
    col_live[0].markdown(
        f'<div class="live-chip"><span class="pulse"></span>'
        f'外部活跃用户：<b>{ext_active} 人</b></div>'
        f'<div style="font-size:11px;color:#8c8c8c;margin-top:4px;padding-left:4px">'
        f'过去 15 分钟内（外部门户）</div>',
        unsafe_allow_html=True)
    col_live[1].markdown(
        f'<div style="font-size:13px;padding-top:6px;color:#595959;">'
        f'📈 外部累计访问：<b>{ext_total:,}</b> 次</div>', unsafe_allow_html=True)
    col_live[2].markdown(
        f'<div style="font-size:13px;padding-top:6px;color:#595959;">'
        f'📅 外部今日访问：<b>{ext_today}</b> 次</div>', unsafe_allow_html=True)
    if st.columns([3,1])[1].button("🔄 刷新监控", use_container_width=True):
        st.rerun()

    st.divider()

    # ── Hourly heatmap ────────────────────────────────────────
    c_left, c_right = st.columns([3,2])
    with c_left:
        st.subheader("分时段访问量（过去 7 天）")
        st.caption("📌 以下图表含模拟历史数据用于演示效果；实时三项指标（活跃/累计/今日）均来自外部门户真实记录。")
        week_ago = datetime.now() - timedelta(days=7)
        week_v = [v for v in all_v_chart if v["ts"] > week_ago]
        if week_v:
            hour_counts = [0]*24
            for v in week_v:
                hour_counts[v["ts"].hour] += 1
            fig_hr = go.Figure()
            fig_hr.add_trace(go.Bar(
                x=list(range(24)), y=hour_counts,
                marker_color=[
                    "#1890ff" if (9<=h<=11 or 14<=h<=16) else "#69b1ff"
                    for h in range(24)],
                name="访问量",
            ))
            fig_hr.add_annotation(text="上午高峰",x=10,y=max(hour_counts)*1.05,
                                   showarrow=False,font=dict(size=11,color="#1890ff"))
            fig_hr.add_annotation(text="下午高峰",x=15,y=max(hour_counts)*1.05,
                                   showarrow=False,font=dict(size=11,color="#1890ff"))
            fig_hr.update_layout(height=240,margin=dict(l=20,r=20,t=30,b=30),
                                  plot_bgcolor="rgba(0,0,0,0)",
                                  xaxis_title="小时",yaxis_title="访问次数",
                                  showlegend=False,
                                  xaxis=dict(tickmode="linear",dtick=2))
            st.plotly_chart(fig_hr, use_container_width=True)

    with c_right:
        st.subheader("页面访问分布")
        page_counts = {}
        for v in all_v_chart: page_counts[v["page"]] = page_counts.get(v["page"],0)+1
        fig_pg = px.pie(values=list(page_counts.values()),
                         names=list(page_counts.keys()), hole=0.5,
                         color_discrete_sequence=px.colors.qualitative.Set2)
        fig_pg.update_traces(textposition="outside",textinfo="percent+label")
        fig_pg.update_layout(height=240,margin=dict(l=0,r=0,t=10,b=10),showlegend=False)
        st.plotly_chart(fig_pg, use_container_width=True)

    # ── Daily trend ───────────────────────────────────────────
    st.subheader("30天日访问量趋势")
    daily = {}
    for v in all_v_chart:
        d = v["ts"].strftime("%Y-%m-%d")
        daily[d] = daily.get(d,0)+1
    dates_sorted = sorted(daily.keys())
    fig_d = go.Figure()
    fig_d.add_trace(go.Scatter(
        x=dates_sorted, y=[daily[d] for d in dates_sorted],
        fill="tozeroy", line=dict(color="#1890ff",width=2),
        fillcolor="rgba(24,144,255,0.12)", name="日访问量",
    ))
    fig_d.update_layout(height=220,margin=dict(l=20,r=20,t=10,b=30),
                         plot_bgcolor="rgba(0,0,0,0)",
                         xaxis_title="",yaxis_title="访问次数",showlegend=False)
    st.plotly_chart(fig_d, use_container_width=True)

    # ── Score distribution + zone pie ────────────────────────
    st.divider()
    sc1, sc2 = st.columns([3,2])
    with sc1:
        st.subheader("AI综合评分分布")
        fig_h=go.Figure()
        fig_h.add_trace(go.Histogram(x=df["AI综合评分"],nbinsx=40,
                                      marker_color="#1890ff",opacity=0.75))
        for x0,x1,col,ann,pos in [(g_thr,100,"#52c41a","绿区·自动通过","top right"),
                                    (r_thr,g_thr,"#faad14","黄区·人工复核","top"),
                                    (0,r_thr,"#ff4d4f","红区·自动淘汰","top left")]:
            fig_h.add_vrect(x0=x0,x1=x1,fillcolor=col,opacity=0.07,line_width=0,
                             annotation_text=ann,annotation_position=pos)
        fig_h.add_vline(x=g_thr,line_dash="dash",line_color="#52c41a",line_width=1.5)
        fig_h.add_vline(x=r_thr,line_dash="dash",line_color="#ff4d4f",line_width=1.5)
        fig_h.update_layout(height=260,margin=dict(l=20,r=20,t=30,b=20),
                             plot_bgcolor="rgba(0,0,0,0)",
                             xaxis_title="AI综合评分",yaxis_title="人数",showlegend=False)
        st.plotly_chart(fig_h, use_container_width=True)
    with sc2:
        st.subheader("评审区间占比")
        fig_p=px.pie(values=[green_n,yellow_n,red_n],
                      names=["绿区（AI通过）","黄区（复核）","红区（淘汰）"],
                      color_discrete_sequence=["#52c41a","#faad14","#ff4d4f"],hole=0.52)
        fig_p.update_traces(textposition="outside",textinfo="percent+label+value")
        fig_p.update_layout(height=260,margin=dict(l=0,r=0,t=10,b=10),showlegend=False)
        st.plotly_chart(fig_p, use_container_width=True)

    # ── Gantt ─────────────────────────────────────────────────
    st.subheader("📅 两周评审流程")
    gantt=[dict(Task="AI解析项目经历与绩效",Start="2025-01-01",Finish="2025-01-04",Phase="第一周·AI初筛"),
           dict(Task="AI批改笔试·生成综合评分",Start="2025-01-04",Finish="2025-01-07",Phase="第一周·AI初筛"),
           dict(Task="区间划分·黄区推送HR",Start="2025-01-07",Finish="2025-01-08",Phase="衔接"),
           dict(Task="HR人工精判黄区（约54人）",Start="2025-01-08",Finish="2025-01-12",Phase="第二周·精判"),
           dict(Task="最终名单审核与公示",Start="2025-01-12",Finish="2025-01-14",Phase="第二周·精判"),
           dict(Task="发展报告推送+48h申诉窗口",Start="2025-01-13",Finish="2025-01-14",Phase="第二周·精判")]
    fig_g=px.timeline(gantt,x_start="Start",x_end="Finish",y="Task",color="Phase",
                       color_discrete_map={"第一周·AI初筛":"#1890ff","衔接":"#722ed1","第二周·精判":"#13c2c2"})
    fig_g.update_yaxes(autorange="reversed",title="")
    fig_g.update_layout(height=240,margin=dict(l=20,r=20,t=10,b=20),
                         plot_bgcolor="rgba(0,0,0,0)",legend=dict(orientation="h",y=-0.3))
    st.plotly_chart(fig_g, use_container_width=True)

# ════════════════════════════════════════════════════════════════
# TAB 2 · 候选人门户（外部独立站点入口）
# ════════════════════════════════════════════════════════════════
with tab2:
    st.markdown("""
    <div style="background:linear-gradient(135deg,#0050b3 0%,#1890ff 60%,#36cfc9 100%);
         color:white;padding:32px 40px;border-radius:16px;margin-bottom:24px;
         box-shadow:0 4px 20px rgba(24,144,255,.25);">
      <h2 style="color:white!important;font-size:22px!important;margin:0 0 10px;">
        🌐 候选人外部门户 — 独立站点
      </h2>
      <p style="margin:0;opacity:.9;font-size:14px;line-height:1.8;">
        候选人门户已部署为 <b>完全独立的外部网站</b>，与本内部系统完全隔离。<br>
        候选人无需登录，直接访问以下地址即可查询报告、提交申诉。
      </p>
    </div>
    """, unsafe_allow_html=True)

    # ── Quick link card ───────────────────────────────────────
    lc1, lc2, lc3 = st.columns([1, 2, 1])
    with lc2:
        st.markdown("""
        <div style="background:white;border:2px solid #1890ff;border-radius:16px;
             padding:36px 32px;text-align:center;
             box-shadow:0 4px 20px rgba(24,144,255,.15);">
          <div style="font-size:52px;margin-bottom:12px;">🔗</div>
          <div style="font-size:22px;font-weight:700;color:#0050b3;
               font-family:monospace;margin-bottom:10px;">
            http://localhost:8504
          </div>
          <div style="font-size:13px;color:#595959;margin-bottom:20px;line-height:1.7;">
            在浏览器新标签页打开上方地址，即可访问候选人门户<br>
            （需先在终端启动：<code>streamlit run candidate_portal.py --server.port 8504</code>）
          </div>
          <div style="display:flex;gap:12px;justify-content:center;flex-wrap:wrap;">
            <span style="background:#e6f4ff;color:#0050b3;padding:6px 16px;
                  border-radius:20px;font-size:13px;font-weight:600;">🔍 查询我的报告</span>
            <span style="background:#fff2f0;color:#cf1322;padding:6px 16px;
                  border-radius:20px;font-size:13px;font-weight:600;">📣 提交申诉</span>
            <span style="background:#f6ffed;color:#389e0d;padding:6px 16px;
                  border-radius:20px;font-size:13px;font-weight:600;">📌 查询申诉进度</span>
          </div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("")
    st.markdown("#### 📊 外部门户实时流量（来自独立站点）")

    # Re-use portal_stats computed in tab1
    ps2 = dm.get_portal_stats(active_minutes=15)
    mc1, mc2, mc3 = st.columns(3)
    mc1.metric("🟢 当前活跃用户", str(ps2["active"]), help="过去15分钟内有访问的外部用户数")
    mc2.metric("📈 累计访问次数", f"{ps2['total']:,}")
    mc3.metric("📅 今日访问次数", str(ps2["today"]))

    st.divider()
    st.markdown("#### 🔄 数据同步说明")
    st.markdown("""
<div class="card card-blue">

**内外数据共享机制**

| 数据类型 | 流向 | 说明 |
|----------|------|------|
| 候选人数据 | 内部 → 外部 | 内部系统启动时自动导出到 `portal_data/candidates.csv` |
| 能力发展报告 | 内部 → 外部 | HR在「能力发展报告」页面生成后自动同步到外部门户 |
| 申诉记录 | 外部 → 内部 | 候选人在外部门户提交后，HR可在「申诉管理」页面看到 |
| 访问统计 | 外部 → 内部 | 外部门户所有访问记录到 `portal_data/analytics.json`，内部监控实时读取 |

</div>
""", unsafe_allow_html=True)

    # (Old embedded portal removed — candidate portal is now a separate app at port 8504)

# ════════════════════════════════════════════════════════════════
# TAB 3 · 简历导入
# ════════════════════════════════════════════════════════════════
with tab3:
    st.subheader("📤 批量简历导入")
    st.markdown("> 支持批量上传 **PDF / DOCX / TXT** 格式简历，AI 自动解析结构化信息并导入评审池。")

    uploaded = st.file_uploader("拖拽上传简历（支持多选）",
                                 type=["pdf","docx","txt"],
                                 accept_multiple_files=True,
                                 label_visibility="collapsed")

    if uploaded and st.button("🤖 开始解析", type="primary"):
        st.session_state.upload_parsed = []
        prog = st.progress(0, text="正在解析…")
        for i, f in enumerate(uploaded):
            text = extract_text(f)
            info = ai_parse(st.session_state.client, text, f.name) \
                   if st.session_state.client and text else {}
            if not info: info = heuristic_parse(text, f.name)
            info["_filename"] = f.name
            info["_source"]   = "Claude AI解析" if st.session_state.client else "规则解析"
            info["_preview"]  = text[:300]
            st.session_state.upload_parsed.append(info)
            prog.progress((i+1)/len(uploaded), text=f"已解析 {i+1}/{len(uploaded)}")
        prog.empty()
        st.success(f"✅ 解析完成！共 {len(uploaded)} 份简历，请在下方确认后导入。")

    if st.session_state.upload_parsed:
        st.divider()
        st.subheader("解析结果预览与编辑")
        confirmed = []
        for i, info in enumerate(st.session_state.upload_parsed):
            with st.expander(f"📄 {info['_filename']}  ·  {info['_source']}", expanded=i==0):
                c1,c2,c3 = st.columns(3)
                name_ = c1.text_input("姓名",     info.get("name",""), key=f"n{i}")
                dept_ = c2.selectbox("当前部门",  DEPTS,
                    index=DEPTS.index(info.get("department","技术部"))
                    if info.get("department") in DEPTS else 0, key=f"d{i}")
                yrs_  = c3.number_input("工龄(年)", 0.5, 20.0,
                    float(info.get("years",3.0)), 0.5, key=f"y{i}")
                s1,s2,s3 = st.columns(3)
                pm_  = s1.slider("岗位匹配度", 0, 100, int(info.get("position_match_estimate",60)),  key=f"pm{i}")
                pe_  = s2.slider("历史绩效",   0, 100, int(info.get("performance_estimate",65)),     key=f"pe{i}")
                wt__ = s3.slider("笔试成绩",   0, 100, int(info.get("written_test_estimate",58)),    key=f"wt{i}")
                hl_ = st.text_input("核心亮点", info.get("highlights",""), key=f"hl{i}")
                st.caption(f"**摘要**：{info.get('_preview','')[:200]}…")
                if st.checkbox(f"✅ 确认导入", key=f"cb{i}"):
                    confirmed.append({"候选人":name_,"员工编号":f"UPL{1000+i}",
                                      "当前部门":dept_,"工龄(年)":yrs_,
                                      "岗位匹配度":float(pm_),"历史绩效":float(pe_),
                                      "笔试成绩":float(wt__),"来源":"简历上传"})
        if confirmed and st.button(f"⬆️ 导入 {len(confirmed)} 名候选人", type="primary"):
            st.session_state.base_df = pd.concat(
                [st.session_state.base_df, pd.DataFrame(confirmed)], ignore_index=True)
            st.session_state.upload_parsed = []
            st.success(f"🎉 成功导入 {len(confirmed)} 名！"); st.rerun()

    st.divider()
    src = st.session_state.base_df["来源"].value_counts().reset_index()
    src.columns = ["来源","人数"]
    st.subheader("已导入数据概况")
    st.dataframe(src, use_container_width=True, hide_index=True)

    # ── 外部自主投递（来自候选人门户）────────────────────────
    st.divider()
    st.subheader("📥 外部自主投递（来自候选人门户）")
    st.markdown("> 候选人通过外部门户（port 8504）自主投递的竞聘申请，HR 审核后可一键导入评审池。")

    _subs = dm.load_submissions()
    if not _subs:
        st.info("暂无外部自主投递记录。候选人可通过 http://localhost:8504 的「在线投递简历」功能提交。")
    else:
        # summary metrics
        sub_total    = len(_subs)
        sub_pending  = sum(1 for s in _subs.values() if s.get("status") == "待审核")
        sub_approved = sum(1 for s in _subs.values() if s.get("status") == "已导入")
        s1, s2, s3 = st.columns(3)
        s1.metric("投递总数", str(sub_total))
        s2.metric("⏳ 待审核", str(sub_pending))
        s3.metric("✅ 已导入", str(sub_approved))

        st.markdown("")
        for sub_id, sub in sorted(_subs.items(), key=lambda x: x[1]["submitted"], reverse=True):
            status_color = {"待审核": "#faad14", "已导入": "#52c41a", "已拒绝": "#ff4d4f"}.get(sub.get("status","待审核"), "#8c8c8c")
            with st.expander(
                f"**{sub_id}** · {sub['name']} · {sub.get('dept','—')} · "
                f"工龄{sub.get('years','?')}年 · "
                f"[{sub.get('status','待审核')}] · {sub['submitted']}",
                expanded=False
            ):
                sc1, sc2 = st.columns([2, 1])
                with sc1:
                    st.markdown(f"""
| 字段 | 内容 |
|------|------|
| 姓名 | {sub['name']} |
| 员工编号 | {sub.get('emp_id','—')} |
| 当前部门 | {sub.get('dept','—')} |
| 工龄 | {sub.get('years','—')} 年 |
| 上传简历 | {sub.get('resume_file','无')} |
| 投递时间 | {sub['submitted']} |
""")
                    if sub.get("statement"):
                        st.markdown(f"**竞聘自述**：{sub['statement'][:200]}{'…' if len(sub.get('statement',''))>200 else ''}")
                with sc2:
                    new_status = st.selectbox(
                        "审核状态",
                        ["待审核", "已导入", "已拒绝"],
                        index=["待审核", "已导入", "已拒绝"].index(sub.get("status","待审核")),
                        key=f"sub_st_{sub_id}"
                    )
                    if st.button("💾 更新状态", key=f"sub_sv_{sub_id}", use_container_width=True):
                        _subs[sub_id]["status"] = new_status
                        dm.save_submissions(_subs)
                        # 如果审核通过，导入评审池
                        if new_status == "已导入":
                            new_row = pd.DataFrame([{
                                "候选人":   sub["name"],
                                "员工编号": sub.get("emp_id", f"EXT{sub_id}"),
                                "当前部门": sub.get("dept", "待确认"),
                                "工龄(年)": float(sub.get("years", 3.0)),
                                "岗位匹配度": float(sub.get("pm_score", 60)),
                                "历史绩效":   float(sub.get("pe_score", 65)),
                                "笔试成绩":   float(sub.get("wt_score", 58)),
                                "来源": "外部投递",
                            }])
                            st.session_state.base_df = pd.concat(
                                [st.session_state.base_df, new_row], ignore_index=True)
                            dm.save_candidates(st.session_state.base_df)
                            st.success(f"✅ {sub['name']} 已导入评审池！")
                        else:
                            st.success("已更新状态")
                        st.rerun()

# ════════════════════════════════════════════════════════════════
# TAB 4 · 员工信息与对比
# ════════════════════════════════════════════════════════════════
with tab4:
    st.subheader("👤 员工信息查询与横向对比")
    mode = st.radio("模式", ["单人详情","多人横向对比"],
                     horizontal=True, label_visibility="collapsed")
    ZICON={"绿区":"🟢","黄区":"🟡","红区":"🔴"}

    if mode=="单人详情":
        sq = st.text_input("🔍 搜索姓名或工号", placeholder="张伟 / EMP10001")
        pool = df[df["候选人"].str.contains(sq)|df["员工编号"].str.contains(sq)]["候选人"].tolist() \
               if sq else ALL_NAMES
        if not pool:
            st.warning("无匹配结果")
        else:
            pick = st.selectbox("选择候选人", pool)
            row  = df[df["候选人"]==pick].iloc[0]
            l,r  = st.columns([1,2])
            with l:
                st.markdown('<div class="card">', unsafe_allow_html=True)
                st.markdown(f"### {row['候选人']}")
                st.markdown(f"`{row['员工编号']}`")
                st.markdown(f"""
| 项目 | 信息 |
|------|------|
| 当前部门 | {row['当前部门']} |
| 工龄 | {row['工龄(年)']} 年 |
| 来源 | {row['来源']} |
| 全体排名 | **第 {row['排名']} 名** / {len(df)} |
| 评审区间 | {ZICON.get(row['评审区间'],'⚪')} **{row['评审区间']}** |
| 状态 | {row['状态']} |
""")
                st.metric("AI综合评分", f"{row['AI综合评分']:.1f}",
                          delta=f"距通过线 {row['AI综合评分']-g_thr:+.1f}分")
                st.markdown('</div>', unsafe_allow_html=True)
                st.markdown("**分项得分**")
                for met, clr, w in [("岗位匹配度","#1890ff",wm),
                                     ("历史绩效",  "#52c41a",wp),
                                     ("笔试成绩",  "#722ed1",wt_)]:
                    sv=row[met]; av=df[met].mean()
                    st.markdown(f"<small>{met}（{w:.0%}）：**{sv:.1f}** · 均值 {av:.1f} "
                                f"{'↑' if sv>av else '↓'}</small>", unsafe_allow_html=True)
                    st.progress(int(sv)/100)
            with r:
                cats=["岗位匹配度","历史绩效","笔试成绩"]
                vals=[row[c] for c in cats]; avgs=[df[c].mean() for c in cats]
                fig_r=go.Figure()
                fig_r.add_trace(go.Scatterpolar(r=vals+[vals[0]],theta=cats+[cats[0]],
                    fill="toself",name=row["候选人"],line_color="#1890ff",
                    fillcolor="rgba(24,144,255,.18)"))
                fig_r.add_trace(go.Scatterpolar(r=avgs+[avgs[0]],theta=cats+[cats[0]],
                    name="全体均值",line_color="#faad14",line_dash="dash"))
                fig_r.update_layout(polar=dict(radialaxis=dict(visible=True,range=[0,100])),
                    height=320,margin=dict(l=40,r=40,t=40,b=30),
                    legend=dict(orientation="h",y=-0.15))
                st.plotly_chart(fig_r, use_container_width=True)
                pct_t=pd.DataFrame({"维度":cats+["综合评分"],
                    "得分":[f"{row[c]:.1f}" for c in cats]+[f"{row['AI综合评分']:.1f}"],
                    "超越全体%":[f"{(df[c]<=row[c]).mean()*100:.0f}%" for c in cats]+
                               [f"{(df['AI综合评分']<=row['AI综合评分']).mean()*100:.0f}%"],
                    "与均值差":[f"{row[c]-df[c].mean():+.1f}" for c in cats]+
                              [f"{row['AI综合评分']-df['AI综合评分'].mean():+.1f}"]})
                st.dataframe(pct_t, use_container_width=True, hide_index=True)
    else:
        st.markdown("选择 **2–4 名**候选人进行横向对比，辅助人工复核决策")
        cnames = st.multiselect("选择候选人（可搜索）", ALL_NAMES,
                                 default=ALL_NAMES[:2], max_selections=4)
        if len(cnames)<2:
            st.info("请至少选择 2 名候选人")
        else:
            crows  = [df[df["候选人"]==n].iloc[0] for n in cnames]
            COLORS = ["#1890ff","#52c41a","#722ed1","#fa8c16"]
            cols = st.columns(len(crows))
            for col,row,color in zip(cols,crows,COLORS):
                with col:
                    z=row["评审区间"]; bc={"绿区":"badge-green","黄区":"badge-yellow","红区":"badge-red"}[z]
                    st.markdown(f'<div class="card" style="border-top:4px solid {color};text-align:center">'
                        f'<h4 style="color:{color};margin:0 0 4px">{row["候选人"]}</h4>'
                        f'<p style="font-size:12px;color:#888;margin:0 0 8px">'
                        f'{row["当前部门"]} · {row["工龄(年)"]}年</p>'
                        f'<p style="font-size:30px;font-weight:700;color:{color};margin:0 0 8px">'
                        f'{row["AI综合评分"]:.1f}</p>'
                        f'<span class="{bc}">{z}</span> '
                        f'<small style="color:#888">第{row["排名"]}名</small>'
                        f'</div>', unsafe_allow_html=True)
            cats=["岗位匹配度","历史绩效","笔试成绩"]
            fig_c=go.Figure()
            for row,color in zip(crows,COLORS):
                v=[row[c] for c in cats]
                fig_c.add_trace(go.Scatterpolar(r=v+[v[0]],theta=cats+[cats[0]],
                    fill="toself",name=row["候选人"],line_color=color,opacity=0.75))
            fig_c.update_layout(polar=dict(radialaxis=dict(visible=True,range=[0,100])),
                height=360,margin=dict(l=40,r=40,t=40,b=30),
                legend=dict(orientation="h",y=-0.15))
            st.plotly_chart(fig_c, use_container_width=True)
            tbl={"指标":cats+["AI综合评分","排名","评审区间","工龄(年)"]}
            for row in crows:
                tbl[row["候选人"]]=[f"{row[c]:.1f}" for c in cats]+[
                    f"{row['AI综合评分']:.1f}",f"第{row['排名']}名",
                    row["评审区间"],str(row["工龄(年)"])]
            st.dataframe(pd.DataFrame(tbl), use_container_width=True, hide_index=True)
            bar_d=[]
            for row in crows:
                for m in cats: bar_d.append({"候选人":row["候选人"],"维度":m,"分数":row[m]})
            fig_b=px.bar(pd.DataFrame(bar_d),x="维度",y="分数",color="候选人",
                barmode="group",color_discrete_sequence=COLORS,range_y=[0,100],height=260)
            fig_b.update_layout(margin=dict(l=20,r=20,t=10,b=20),
                plot_bgcolor="rgba(0,0,0,0)",legend=dict(orientation="h",y=1.1))
            st.plotly_chart(fig_b, use_container_width=True)
            best=max(crows,key=lambda x:x["AI综合评分"])
            st.info(f"**AI辅助建议**：当前权重下 **{best['候选人']}** 综合得分最高"
                    f"（{best['AI综合评分']:.1f}分），如仅有 1 个名额建议优先通过。")

# ════════════════════════════════════════════════════════════════
# TAB 5 · 人工复核
# ════════════════════════════════════════════════════════════════
with tab5:
    ydf  = df[df["评审区间"]=="黄区"].copy()
    done = len(st.session_state.human_reviews)
    h1,h2,h3 = st.columns(3)
    h1.metric("黄区总人数",str(yellow_n)); h2.metric("已完成",str(done)); h3.metric("待复核",str(yellow_n-done))
    st.progress(done/yellow_n if yellow_n>0 else 0,
                text=f"复核进度：{done}/{yellow_n}（{done/yellow_n:.0%}）" if yellow_n>0 else "")
    st.divider()
    pending   = ydf[~ydf["员工编号"].isin(st.session_state.human_reviews)]
    completed = ydf[ ydf["员工编号"].isin(st.session_state.human_reviews)]
    if len(pending)>0:
        rname = st.selectbox("选择待复核候选人", pending["候选人"].tolist())
        rev   = ydf[ydf["候选人"]==rname].iloc[0]
        l,r   = st.columns(2)
        with l:
            st.markdown(f"#### {rev['候选人']} · 评分详情")
            for met,w,clr in [("岗位匹配度",wm,"#1890ff"),("历史绩效",wp,"#52c41a"),("笔试成绩",wt_,"#722ed1")]:
                sc=rev[met]; av=df[met].mean()
                st.markdown(f"**{met}**（{w:.0%}）— {sc:.1f}分 "
                            f"{'↑' if sc>av else '↓'} 均值 {av:.1f}")
                st.progress(int(sc)/100)
            st.metric("AI综合评分",f"{rev['AI综合评分']:.1f}分",
                      delta=f"距通过线 {rev['AI综合评分']-g_thr:+.1f}分",
                      delta_color="normal" if rev['AI综合评分']>=g_thr else "inverse")
        with r:
            st.markdown("#### HR 复核决策")
            st.info(f"评分 **{rev['AI综合评分']:.1f}分**，处于临界区间（{r_thr}–{g_thr}分）。")
            notes = st.text_area("复核备注（必填）",
                                  placeholder="如：面谈后确认候选人有丰富B端PM经验…",
                                  height=110, key=f"rev_{rev['员工编号']}")
            b1,b2 = st.columns(2)
            with b1:
                if st.button("✅ 通过",type="primary",use_container_width=True,disabled=not notes):
                    st.session_state.human_reviews[rev["员工编号"]]={"decision":"通过","notes":notes,"candidate":rev["候选人"]}
                    st.success(f"已确认 {rev['候选人']} 通过！"); st.rerun()
            with b2:
                if st.button("❌ 淘汰",use_container_width=True,disabled=not notes):
                    st.session_state.human_reviews[rev["员工编号"]]={"decision":"淘汰","notes":notes,"candidate":rev["候选人"]}
                    st.warning(f"已确认 {rev['候选人']} 淘汰。"); st.rerun()
    else:
        st.success("🎉 所有黄区候选人已完成人工复核！")
    if len(completed)>0:
        st.divider(); st.subheader("已完成复核记录")
        recs=[]
        for _,row in completed.iterrows():
            rv=st.session_state.human_reviews.get(row["员工编号"])
            if rv: recs.append({"候选人":row["候选人"],"AI评分":row["AI综合评分"],
                                 "HR决策":rv["decision"],"备注":rv["notes"][:40]+"…" if len(rv["notes"])>40 else rv["notes"]})
        st.dataframe(pd.DataFrame(recs), use_container_width=True, hide_index=True)

# ════════════════════════════════════════════════════════════════
# TAB 6 · 能力发展报告
# ════════════════════════════════════════════════════════════════
with tab6:
    st.subheader("📋 能力发展报告")
    st.markdown("> 个性化报告包含 **分项得分分析**、**学习超链接** 及 **行动计划**，支持 PDF / Word 双格式下载。")
    rpt_pool = df[df["评审区间"].isin(["红区","黄区"])].head(80)
    cp, cr = st.columns([1,2])
    with cp:
        sq = st.text_input("搜索候选人", placeholder="输入姓名…", key="rpt_sq")
        pool2 = rpt_pool[rpt_pool["候选人"].str.contains(sq)]["候选人"].tolist() if sq else rpt_pool["候选人"].tolist()
        rname = st.selectbox("选择候选人（落榜/待复核）", pool2, key="rpt_name")
        rc = df[df["候选人"]==rname].iloc[0]
        zi={"绿区":"🟢","黄区":"🟡","红区":"🔴"}.get(rc["评审区间"],"⚪")
        st.markdown(f"**{rc['候选人']}** {zi} {rc['评审区间']}")
        st.caption(f"{rc['当前部门']} · 工龄 {rc['工龄(年)']}年")
        for met,clr in [("岗位匹配度","#1890ff"),("历史绩效","#52c41a"),("笔试成绩","#722ed1")]:
            sc=rc[met]; av=df[met].mean()
            st.markdown(f"<small>{met}：**{sc:.1f}** / 均值 {av:.1f}</small>",unsafe_allow_html=True)
            st.progress(int(sc)/100)
        gen = st.button("🤖 生成能力发展报告",type="primary",use_container_width=True)
    with cr:
        eid = rc["员工编号"]
        if gen:
            with st.spinner("正在生成个性化报告…"):
                rpt = build_report(rc, df, g_thr, wm, wp, wt_)
                if st.session_state.client:
                    try:
                        resp = st.session_state.client.messages.create(
                            model="claude-sonnet-4-6", max_tokens=2000,
                            system="你是腾讯人才发展顾问，报告温和、专业、激励性强，保留所有链接和数据。",
                            messages=[{"role":"user","content":
                                f"请用更流畅自然的语气改写以下报告，保留所有数据与链接，字数控制在700字以内：\n\n{rpt}"}])
                        rpt = resp.content[0].text
                    except: pass
                st.session_state.generated_reports[eid] = rpt
                # 同步到外部门户共享文件，候选人可在外部门户查看
                _rpts = dm.load_reports()
                _rpts[eid] = rpt
                dm.save_reports(_rpts)
        if eid in st.session_state.generated_reports:
            rpt_text = st.session_state.generated_reports[eid]
            st.markdown('<div class="report-box">',unsafe_allow_html=True)
            st.markdown(rpt_text)
            st.markdown('</div>',unsafe_allow_html=True)
            d1,d2 = st.columns(2)
            with d1:
                wb = to_word_bytes(rpt_text)
                if wb: st.download_button("📄 下载 Word (.docx)",data=wb,
                    file_name=f"{rname}_能力发展报告.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            with d2:
                pb = to_pdf_bytes(rpt_text)
                if pb: st.download_button("📑 下载 PDF (.pdf)",data=pb,
                    file_name=f"{rname}_能力发展报告.pdf",
                    mime="application/pdf",use_container_width=True)
        else:
            st.markdown('<div style="text-align:center;padding:60px 20px;color:#8c8c8c;'
                'border:2px dashed #d9d9d9;border-radius:12px">'
                '<div style="font-size:48px">📋</div>'
                '<div style="margin-top:12px;font-size:15px">点击左侧按钮生成报告</div>'
                '<div style="margin-top:8px;font-size:12px;color:#bfbfbf">'
                '含学习超链接 · 支持 PDF / Word 双格式下载</div>'
                '</div>', unsafe_allow_html=True)

# ════════════════════════════════════════════════════════════════
# TAB 7 · 申诉管理中心
# ════════════════════════════════════════════════════════════════
with tab7:
    st.subheader("📣 申诉管理中心")

    # 合并外部门户申诉（portal_data/appeals.json）+ 内部录入申诉
    _ext_ap = dm.load_appeals()
    _merged = {**_ext_ap, **st.session_state.appeals}

    if not _merged:
        st.info("暂无申诉记录。候选人可通过外部门户（http://localhost:8504）提交申诉，HR也可在下方代为录入。")
    else:
        aps = _merged
        # ── Summary metrics ───────────────────────────────────
        total_a  = len(aps)
        passed_a = sum(1 for a in aps.values() if "通过" in a["status"])
        pending_a= sum(1 for a in aps.values() if "待审核" in a["status"] or "审核中" in a["status"])
        high_urg = sum(1 for a in aps.values() if a.get("urgency")=="高")
        am1,am2,am3,am4 = st.columns(4)
        am1.metric("申诉总数",str(total_a))
        am2.metric("⚡ 高优先级",str(high_urg))
        am3.metric("✅ 通过申诉",str(passed_a))
        am4.metric("⏳ 待处理",  str(pending_a))

        st.divider()

        # ── Visualizations ────────────────────────────────────
        vc1, vc2, vc3 = st.columns(3)

        with vc1:
            st.markdown("**申诉类型分布**")
            cat_c = {}
            for a in aps.values(): cat_c[a.get("ai_category",a["category"])] = cat_c.get(a.get("ai_category",a["category"]),0)+1
            fig_ac = px.pie(values=list(cat_c.values()), names=list(cat_c.keys()),
                             hole=0.5, color_discrete_sequence=px.colors.qualitative.Set2)
            fig_ac.update_traces(textposition="outside", textinfo="percent+label")
            fig_ac.update_layout(height=220, margin=dict(l=0,r=0,t=10,b=10), showlegend=False)
            st.plotly_chart(fig_ac, use_container_width=True)

        with vc2:
            st.markdown("**申诉状态分布**")
            st_c = {}
            for a in aps.values(): st_c[a["status"]] = st_c.get(a["status"],0)+1
            fig_st = px.bar(x=list(st_c.keys()), y=list(st_c.values()),
                             color=list(st_c.keys()),
                             color_discrete_sequence=["#faad14","#1890ff","#52c41a","#ff4d4f"])
            fig_st.update_layout(height=220, margin=dict(l=20,r=10,t=10,b=20),
                                  plot_bgcolor="rgba(0,0,0,0)", showlegend=False,
                                  xaxis_title="", yaxis_title="数量")
            st.plotly_chart(fig_st, use_container_width=True)

        with vc3:
            st.markdown("**优先级分布**")
            urg_c = {}
            for a in aps.values(): urg_c[a.get("urgency","中")] = urg_c.get(a.get("urgency","中"),0)+1
            fig_ug = px.bar(x=list(urg_c.keys()), y=list(urg_c.values()),
                             color=list(urg_c.keys()),
                             color_discrete_map={"高":"#ff4d4f","中":"#faad14","低":"#52c41a"})
            fig_ug.update_layout(height=220, margin=dict(l=20,r=10,t=10,b=20),
                                  plot_bgcolor="rgba(0,0,0,0)", showlegend=False,
                                  xaxis_title="优先级", yaxis_title="数量")
            st.plotly_chart(fig_ug, use_container_width=True)

        # ── Submission timeline ───────────────────────────────
        st.markdown("**申诉提交时间线**")
        times = [a["submitted"] for a in aps.values()]
        names_a = [a["candidate"] for a in aps.values()]
        cats_a  = [a.get("ai_category", a["category"]) for a in aps.values()]
        urgs_a  = [a.get("urgency","中") for a in aps.values()]
        fig_tl = go.Figure()
        clr_map = {"高":"#ff4d4f","中":"#faad14","低":"#52c41a"}
        for i,(t,n,c,u) in enumerate(zip(times,names_a,cats_a,urgs_a)):
            fig_tl.add_trace(go.Scatter(
                x=[t], y=[i], mode="markers+text",
                marker=dict(size=14, color=clr_map.get(u,"#1890ff"),
                            symbol="circle", line=dict(width=2, color="white")),
                text=[f"{n}·{c}"], textposition="middle right",
                name=u, showlegend=False,
            ))
        fig_tl.update_layout(height=max(120, len(aps)*50), margin=dict(l=20,r=200,t=20,b=20),
                              plot_bgcolor="rgba(0,0,0,0)",
                              yaxis=dict(showticklabels=False, showgrid=False),
                              xaxis=dict(showgrid=True, gridcolor="#f0f0f0"))
        st.plotly_chart(fig_tl, use_container_width=True)

        # ── Detail management ─────────────────────────────────
        st.divider()
        st.subheader("逐条处理")
        sort_by = st.radio("排序方式", ["⚡ 高优先级优先","📅 时间倒序"],
                            horizontal=True, label_visibility="collapsed")
        items = sorted(aps.items(),
                       key=lambda x: ({"高":0,"中":1,"低":2}.get(x[1].get("urgency","中"),1), x[1]["submitted"])
                       if "优先" in sort_by else lambda x: x[1]["submitted"])
        # Ensure items is a list after sort
        items = list(aps.items())
        if "优先" in sort_by:
            items.sort(key=lambda x: ({"高":0,"中":1,"低":2}.get(x[1].get("urgency","中"),1)))
        else:
            items.sort(key=lambda x: x[1]["submitted"], reverse=True)

        for aid, ap in items:
            urg_color = {"高":"#ff4d4f","中":"#faad14","低":"#52c41a"}.get(ap.get("urgency","中"),"#1890ff")
            with st.expander(
                f"**{aid}** · {ap['candidate']} · "
                f"[{ap.get('ai_category', ap['category'])}] · "
                f"{ap['status']} · ⚡{ap.get('urgency','中')} · {ap['submitted']}",
                expanded=False
            ):
                e1, e2 = st.columns([3,2])
                with e1:
                    st.markdown(f"""
| 字段 | 内容 |
|------|------|
| 申诉人 | {ap['candidate']}（{ap['emp_id']}） |
| 原始评分 | {ap['score']:.1f} 分 |
| AI分类 | **{ap.get('ai_category', ap['category'])}** |
| 优先级 | <span style="color:{urg_color};font-weight:600">{ap.get('urgency','中')}</span> |
| 联系方式 | {ap['contact']} |
| 附件 | {', '.join(ap['files']) if ap['files'] else '无'} |
""", unsafe_allow_html=True)
                    st.markdown(f"**申诉内容**：{ap['reason']}")
                    if ap.get("key_claims"):
                        st.markdown(f"**AI提炼诉求**：" + "；".join(ap["key_claims"]))
                    if ap.get("auto_response"):
                        st.info(f"**AI初步回复**：{ap['auto_response']}")
                with e2:
                    new_status = st.selectbox("更新状态",
                        ["⏳ 待审核","🔍 审核中","✅ 通过申诉","❌ 维持原判"],
                        index=["⏳ 待审核","🔍 审核中","✅ 通过申诉","❌ 维持原判"].index(ap["status"])
                        if ap["status"] in ["⏳ 待审核","🔍 审核中","✅ 通过申诉","❌ 维持原判"] else 0,
                        key=f"st_{aid}")
                    hr_cmt = st.text_area("HR回复意见", ap.get("hr_comment",""),
                                           height=90, key=f"cmt_{aid}")
                    if st.button("💾 保存", key=f"save_{aid}", use_container_width=True):
                        # 写入 session_state（内部录入的）或外部文件（外部提交的）
                        if aid in st.session_state.appeals:
                            st.session_state.appeals[aid]["status"]     = new_status
                            st.session_state.appeals[aid]["hr_comment"] = hr_cmt
                        # 无论来源，同步回外部共享文件
                        _all_ap = dm.load_appeals()
                        if aid in _all_ap:
                            _all_ap[aid]["status"]     = new_status
                            _all_ap[aid]["hr_comment"] = hr_cmt
                            dm.save_appeals(_all_ap)
                        st.success("已更新"); st.rerun()

    # ── Quick submit (HR side) ────────────────────────────────
    st.divider()
    with st.expander("➕ HR代为录入申诉（快速通道）"):
        qa1,qa2=st.columns(2)
        with qa1:
            qname=st.selectbox("候选人", ALL_NAMES, key="hr_ap_name")
            qrow =df[df["候选人"]==qname].iloc[0]
        with qa2:
            qcat=st.selectbox("申诉类型",APPEAL_CATS,key="hr_ap_cat")
        qreason=st.text_area("申诉内容",height=80,key="hr_ap_reason")
        qcontact=st.text_input("联系方式",key="hr_ap_contact")
        if st.button("录入",key="hr_ap_submit",disabled=not (qreason and qcontact)):
            cls=classify_appeal_rules(qreason)
            _existing = dm.load_appeals()
            aid=f"AP{len(_existing)+len(st.session_state.appeals)+1001}"
            _new_ap={
                "candidate":qname,"emp_id":qrow["员工编号"],"score":qrow["AI综合评分"],
                "category":qcat,"reason":qreason,"contact":qcontact,"files":[],
                "status":"⏳ 待审核","submitted":datetime.now().strftime("%Y-%m-%d %H:%M"),
                "hr_comment":"","ai_category":cls["category"],"urgency":cls["urgency"],
                "key_claims":cls["key_claims"],"auto_response":cls["auto_response"],
                "source": "HR代录"
            }
            st.session_state.appeals[aid] = _new_ap
            # 同步写入外部共享文件，使外部门户可查询进度
            _existing[aid] = _new_ap
            dm.save_appeals(_existing)
            st.success(f"已录入，申诉编号 {aid}"); st.rerun()

# ════════════════════════════════════════════════════════════════
# TAB 8 · 效果评估（动态）
# ════════════════════════════════════════════════════════════════
with tab8:
    st.subheader("📈 效果评估（动态数据看板）")

    kpi_df = st.session_state.kpi_df.copy()

    # ── Live KPI gauges ───────────────────────────────────────
    st.markdown("#### 最新一期核心指标 vs 董事会目标")

    latest = kpi_df.iloc[-1]
    TARGETS = {"评审时长_AI后(天)":14, "人工复核率_AI后(%)":15,
               "投诉率_AI后(%)":5,    "1年离职率_AI后(%)":10}
    LABELS  = {"评审时长_AI后(天)":"评审时长",  "人工复核率_AI后(%)":"人工复核率",
               "投诉率_AI后(%)":"预估投诉率",  "1年离职率_AI后(%)":"1年内离职率"}
    UNITS   = {"评审时长_AI后(天)":"天",  "人工复核率_AI后(%)":"%",
               "投诉率_AI后(%)":"%",    "1年离职率_AI后(%)":"%"}

    gcols = st.columns(4)
    for col, (k, tgt) in zip(gcols, TARGETS.items()):
        with col:
            val = float(latest[k])
            good = val < tgt
            clr  = "#52c41a" if good else "#ff4d4f"
            fig_g = go.Figure(go.Indicator(
                mode="gauge+number+delta", value=val,
                number={"suffix": UNITS[k], "font":{"size":26,"color":clr}},
                delta={"reference":tgt,"valueformat":".1f",
                       "decreasing":{"color":"#52c41a"},"increasing":{"color":"#ff4d4f"}},
                gauge={"axis":{"range":[0,tgt*1.5]},
                       "bar":{"color":clr},
                       "steps":[{"range":[0,tgt*.7],"color":"#f6ffed"},
                                 {"range":[tgt*.7,tgt],"color":"#fffbe6"},
                                 {"range":[tgt,tgt*1.5],"color":"#fff2f0"}],
                       "threshold":{"line":{"color":"#ff4d4f","width":2},"value":tgt}},
                title={"text":f"{LABELS[k]}<br>"
                              f"<span style='font-size:11px;color:gray'>目标≤{tgt}{UNITS[k]}</span>"},
            ))
            fig_g.update_layout(height=210, margin=dict(l=20,r=20,t=55,b=10))
            st.plotly_chart(fig_g, use_container_width=True)

    st.divider()

    # ── Before / After comparison ─────────────────────────────
    st.markdown("#### AI赋能前后趋势对比")
    col_sel1, col_sel2 = st.columns([3,1])
    with col_sel1:
        metric_opt = st.selectbox("选择对比指标", [
            "评审时长（天）",
            "投诉率（%）",
            "人工复核率（%）",
            "1年离职率（%）",
        ])
    with col_sel2:
        chart_type = st.radio("图表类型", ["折线","面积"], horizontal=True,
                               label_visibility="collapsed")

    col_map = {
        "评审时长（天）":  ("评审时长_AI前(天)",  "评审时长_AI后(天)"),
        "投诉率（%）":     ("投诉率_AI前(%)",     "投诉率_AI后(%)"),
        "人工复核率（%）": (None,                 "人工复核率_AI后(%)"),
        "1年离职率（%）":  (None,                 "1年离职率_AI后(%)"),
    }
    before_col, after_col = col_map[metric_opt]

    fig_ba = go.Figure()
    fill_mode = "tozeroy" if chart_type=="面积" else "none"
    if before_col:
        fig_ba.add_trace(go.Scatter(
            x=kpi_df["日期"], y=kpi_df[before_col], name="引入AI前",
            line=dict(color="#ff7875",width=2,dash="dot"), fill=fill_mode,
            fillcolor="rgba(255,120,117,.1)"))
    fig_ba.add_trace(go.Scatter(
        x=kpi_df["日期"], y=kpi_df[after_col], name="引入AI后",
        line=dict(color="#1890ff",width=2.5), fill=fill_mode,
        fillcolor="rgba(24,144,255,.1)"))
    target_val = list(TARGETS.values())[["评审时长（天）","投诉率（%）","人工复核率（%）","1年离职率（%）"].index(metric_opt)]
    fig_ba.add_hline(y=target_val, line_dash="dash", line_color="#52c41a",
                      annotation_text=f"董事会目标 {target_val}",
                      annotation_position="right")
    fig_ba.update_layout(height=300, margin=dict(l=20,r=80,t=20,b=30),
                          plot_bgcolor="rgba(0,0,0,0)",
                          yaxis_title=metric_opt,
                          legend=dict(orientation="h",y=1.1))
    st.plotly_chart(fig_ba, use_container_width=True)

    # ── Editable data table ───────────────────────────────────
    st.divider()
    st.markdown("#### 📝 数据管理（可直接编辑或新增记录）")
    st.caption("直接在表格中修改数值，或在最后一行新增数据，系统将自动更新上方图表。")

    edited_df = st.data_editor(
        st.session_state.kpi_df,
        use_container_width=True,
        num_rows="dynamic",
        hide_index=True,
        column_config={
            "日期":              st.column_config.TextColumn("日期", width="medium"),
            "评审时长_AI前(天)":  st.column_config.NumberColumn("评审时长_AI前(天)",  min_value=1,  max_value=60,  step=1),
            "评审时长_AI后(天)":  st.column_config.NumberColumn("评审时长_AI后(天)",  min_value=1,  max_value=30,  step=0.5),
            "投诉率_AI前(%)":     st.column_config.NumberColumn("投诉率_AI前(%)",     min_value=0,  max_value=50,  step=0.1),
            "投诉率_AI后(%)":     st.column_config.NumberColumn("投诉率_AI后(%)",     min_value=0,  max_value=20,  step=0.1),
            "人工复核率_AI后(%)": st.column_config.NumberColumn("人工复核率_AI后(%)", min_value=0,  max_value=30,  step=0.1),
            "1年离职率_AI后(%)":  st.column_config.NumberColumn("1年离职率_AI后(%)",  min_value=0,  max_value=30,  step=0.1),
        },
        key="kpi_editor",
    )

    if st.button("💾 保存数据更新", type="primary"):
        st.session_state.kpi_df = edited_df
        st.success("✅ 数据已更新，图表已同步刷新！"); st.rerun()

    # ── Summary comparison stats ──────────────────────────────
    st.divider()
    st.markdown("#### 引入AI前后 · 均值对比汇总")
    if "评审时长_AI前(天)" in kpi_df.columns:
        summary = pd.DataFrame({
            "指标":     ["评审时长(天)","投诉率(%)","人工复核率(%)","1年离职率(%)"],
            "AI引入前(均值)": [
                f"{kpi_df['评审时长_AI前(天)'].mean():.1f}",
                f"{kpi_df['投诉率_AI前(%)'].mean():.1f}",
                "—",
                "—",
            ],
            "AI引入后(均值)": [
                f"{kpi_df['评审时长_AI后(天)'].mean():.1f}",
                f"{kpi_df['投诉率_AI后(%)'].mean():.1f}",
                f"{kpi_df['人工复核率_AI后(%)'].mean():.1f}",
                f"{kpi_df['1年离职率_AI后(%)'].mean():.1f}",
            ],
            "董事会目标": ["≤14天","≤5%","≤15%","≤10%"],
            "改善幅度": [
                f"↓{kpi_df['评审时长_AI前(天)'].mean()-kpi_df['评审时长_AI后(天)'].mean():.1f}天",
                f"↓{kpi_df['投诉率_AI前(%)'].mean()-kpi_df['投诉率_AI后(%)'].mean():.1f}%",
                "新指标",
                "新指标",
            ],
        })
        st.dataframe(summary, use_container_width=True, hide_index=True)
